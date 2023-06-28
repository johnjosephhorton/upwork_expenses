from datetime import datetime
from collections.abc import Sequence
from collections import Counter
import requests
import tempfile

import os 
import copy

import pandas as pd
from docx import Document
from PyPDF2 import PdfMerger

from dotenv import load_dotenv
load_dotenv()

INVOICE_PATH = os.path.join(os.getcwd(), "upwork_invoices")
DEFAULT_FUNDING_SOURCE = os.getenv("DEFAULT_FUNDING_SOURCE")
SHEETS_URL = os.getenv("SHEETS_URL")

tempfile_excel = tempfile.NamedTemporaryFile()

def download_google_sheet_as_excel(url):
    response = requests.get(url)
    with open(tempfile_excel.name, 'wb') as f:
        f.write(response.content)

download_google_sheet_as_excel(SHEETS_URL)

all_sheets = pd.read_excel(tempfile_excel.name, sheet_name=None)

freelancers_data = all_sheets['freelancers']
dict_list = {x["Freelancer"]:x for x in freelancers_data.to_dict('records')}

def round(num, digits = 2):
    return float(f"{num:.{digits}f}")

def format_as_money(num, digits = 2):
    return f"${round(num, digits):,.{digits}f}"

class Transaction:
    def __init__(self, dict):
        for key, value in dict.items():
            setattr(self, key, value)
        
        freelancer = self.__dict__['Freelancer']
        if freelancer not in dict_list:
            self.funding_source = DEFAULT_FUNDING_SOURCE
        else:
            self.funding_source = dict_list[self.__dict__['Freelancer']]['Funding source']

    def __repr__(self):
        return f"Expense({self.__dict__})"
    
    def get_upwork_invoice(self):
        path = os.path.join(INVOICE_PATH, "T" + str(self.__dict__['Ref ID']) + ".pdf")
        return path
    
    def summary_dict(self, fields = ["Ref ID", "Date", "Description", "Amount"]):
        return {field: self.__dict__[field] for field in fields}       
    
class Transactions(Sequence):
    def __init__(self, transactions):
        self._transactions = list(transactions)

    def __getitem__(self, index):
        return self._transactions[index]
        
    def __len__(self):
        return len(self._transactions)
    
    def filter_by_funding_source(self, funding_source):
        self._transactions = [expense for expense in self._transactions if expense.funding_source == funding_source]
        return None

    def filter_out_processed(self):
        self._transactions = [expense for expense in self._transactions if not expense.Processed]
        return None

    def filter_out_type(self, type):
        self._transactions = [expense for expense in self._transactions if expense.Type != type]
        return None
    
    def filter_by_date(self, start, end):
        self._transactions = [expense for expense in self._transactions if expense.Date >= start and expense.Date <= end]
        return None
    
    def filter_by_freelancer(self, freelancer):
        self._transactions = [expense for expense in self._transactions if expense.Freelancer == freelancer]
        return None 
    
    def total_charges(self):
        return sum(expense.Amount for expense in self._transactions if expense.Amount < 0)
    
    def total_credits(self):
        return sum(expense.Amount for expense in self._transactions if expense.Amount > 0)
    
    def summary_by_type(self):
        d = dict({})
        total = 0
        for expense in self._transactions:
            total += expense.Amount
            if expense.Type not in d:
                d[expense.Type] = expense.Amount
            else:
                d[expense.Type] += expense.Amount
        d['Total'] = total
        for key in d:
            d[key] = round(d[key])
        return d
    
    def combine_pdfs(self, output_filename = "combined.pdf"):
        pdf_list = [expense.get_upwork_invoice() for expense in self._transactions]
        merger = PdfMerger()

        for pdf in pdf_list:
            merger.append(pdf)

        merger.write(output_filename)
        merger.close()

    
    @property
    def unique_freelancers(self):
        freelancers = set()
        for expense in self._transactions:
            if hasattr(expense, 'Freelancer'):
                if isinstance(expense.Freelancer, str): 
                    freelancers.add(expense.Freelancer)
        return freelancers
    
    @property
    def charges_types(self):
        entry_types = [t.Type for t in self._transactions]
        return Counter(entry_types)

    @property
    def charges_per_freelancer(self):
        charges_dict = dict({})
        freelancers = self.unique_freelancers
        for expense in self._transactions:
            if hasattr(expense, 'Freelancer'):
                if isinstance(expense.Freelancer, str): 
                    if expense.Freelancer not in charges_dict:
                        charges_dict[expense.Freelancer] = expense.Amount
                    else:
                        charges_dict[expense.Freelancer] += expense.Amount
        return charges_dict

def add_table_from_dict(doc, dict):
    num_rows = len(dict.keys())
    table = doc.add_table(rows=num_rows, cols=2)

    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"{list(dict.items())[i][j]}"

def add_summary(doc, summary):
    num_rows = len(summary.keys())
    table = doc.add_table(rows=num_rows, cols=2)

    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"{list(summary.items())[i][j]}"

def create_invoice(transactions, invoice_name, invoice_from, invoice_to):
    invoice_date = datetime.now().strftime('%Y_%m_%d')
    doc = Document()
    doc.add_paragraph(f"From: {invoice_from}\nTo:{invoice_to}\nDate: {datetime.now().strftime('%b %d, %Y')}")
    doc.add_heading('Total', level = 1)
    doc.add_paragraph(f"Total amount for re-imbursement: {format_as_money(transactions.total_charges())}")
    doc.add_paragraph(f"Total invoices:{len(transactions)}")
    doc.add_heading('Charges by type', level = 2)
    add_summary(doc, transactions.summary_by_type())
    doc.add_heading('Summary by Freelancer', level = 1)
    freelancers = transactions.unique_freelancers
    for freelancer in freelancers:
        total = 0
        doc.add_heading(f"{freelancer}", level = 2)
        add_table_from_dict(doc, dict_list[freelancer])
        doc.add_heading("By type of charge", level = 3)
        expenses = copy.deepcopy(transactions)
        expenses.filter_by_freelancer(freelancer)
        summary = expenses.summary_by_type()
        total += sum(summary.values())
        add_table_from_dict(doc, summary)
        doc.add_heading(f"Detailed transactions for {freelancer}", level = 3)
        doc.add_paragraph("See invoices by Ref_ID for details")
        for transaction in expenses:
            add_table_from_dict(doc, transaction.summary_dict())
            doc.add_paragraph("")
    doc.save(f"{invoice_name}_{invoice_date}.docx")
    transactions.combine_pdfs(output_filename = f"invoices_{invoice_name}_{invoice_date}.pdf")

